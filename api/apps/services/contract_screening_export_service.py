#
#  Copyright 2026 The InfiniFlow Authors. All Rights Reserved.
#
#  Licensed under the Apache License, Version 2.0 (the "License");
#  you may not use this file except in compliance with the License.
#  You may obtain a copy of the License at
#
#      http://www.apache.org/licenses/LICENSE-2.0
#
#  Unless required by applicable law or agreed to in writing, software
#  distributed under the License is distributed on an "AS IS" BASIS,
#  WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
#  See the License for the specific language governing permissions and
#  limitations under the License.
#
from __future__ import annotations

from pathlib import Path
from typing import Any

import openpyxl
from docx import Document

from api.apps.services.contract_screening_service import ContractScreeningError
from api.db.services import contract_screening_service as contract_screening_db_service
from common.misc_utils import get_uuid


EXPORT_FORMATS = {"excel": ".xlsx", "word": ".docx"}


def create_screening_export(
    *,
    tenant_id: str,
    user_id: str,
    task_id: str,
    export_format: str,
    output_dir: str | Path,
    repository: Any = contract_screening_db_service,
) -> dict[str, Any]:
    export_format = str(export_format or "").lower()
    if export_format not in EXPORT_FORMATS:
        raise ContractScreeningError("Unsupported export format")

    payload = repository.build_results_payload(tenant_id, task_id, user_id=user_id)
    if not payload:
        raise ContractScreeningError("Task not found")
    if payload.get("status") != "done":
        raise ContractScreeningError("Only completed screening tasks can be exported")

    export_id = get_uuid()
    suffix = EXPORT_FORMATS[export_format]
    file_name = f"contract-screening-{task_id}-{export_id}{suffix}"
    output_path = Path(output_dir)
    output_path.mkdir(parents=True, exist_ok=True)
    file_path = output_path / file_name

    _create_export_record(repository, {
        "id": export_id,
        "task_id": task_id,
        "tenant_id": tenant_id,
        "user_id": user_id,
        "format": export_format,
        "status": "running",
        "file_name": file_name,
        "file_key": "",
        "error": "",
    })

    if export_format == "excel":
        _write_excel(payload, file_path)
    else:
        _write_word(payload, file_path)

    _mark_export_done(repository, export_id, file_name, str(file_path))
    return {
        "export_id": export_id,
        "status": "done",
        "file_name": file_name,
        "file_key": str(file_path),
    }


def _create_export_record(repository: Any, record: dict[str, Any]) -> None:
    create = getattr(repository, "create_export_record", None)
    if callable(create):
        create(record)
        return
    repository.ContractScreeningExportService.insert(**record)


def _mark_export_done(repository: Any, export_id: str, file_name: str, file_key: str) -> None:
    mark = getattr(repository, "mark_export_done", None)
    if callable(mark):
        mark(export_id, file_name, file_key)
        return
    repository.ContractScreeningExportService.update_by_id(export_id, {
        "status": "done",
        "file_name": file_name,
        "file_key": file_key,
        "error": "",
    })


def _write_excel(payload: dict[str, Any], file_path: Path) -> None:
    workbook = openpyxl.Workbook()
    summary = workbook.active
    summary.title = "筛选摘要"
    summary.append(["Prompt", payload.get("prompt", "")])
    summary.append(["任务状态", payload.get("status", "")])
    summary.append(["命中数量", len(payload.get("items") or [])])
    summary.append(["跳过文档", str(payload.get("skipped") or {})])
    summary.append(["筛选条件", _conditions_text(payload)])

    results = workbook.create_sheet("合同结果")
    results.append(["合同名称", "文档 id", "命中状态", "风险", "分数", "原因", "建议动作"])
    for item in payload.get("items") or []:
        results.append([
            item.get("name") or item.get("title") or "",
            item.get("contract_id") or item.get("document_id") or "",
            item.get("status") or "",
            item.get("risk") or (item.get("meta") or {}).get("risk", ""),
            item.get("score") or (item.get("meta") or {}).get("score", ""),
            item.get("reason") or "",
            "；".join(item.get("actions") or []),
        ])

    evidence_sheet = workbook.create_sheet("证据明细")
    evidence_sheet.append(["合同名称", "文档 id", "条件 id", "页码", "Chunk id", "证据文本"])
    for item in payload.get("items") or []:
        for evidence in item.get("evidence") or []:
            evidence_sheet.append([
                item.get("name") or item.get("title") or "",
                item.get("contract_id") or item.get("document_id") or "",
                evidence.get("condition_id") or "",
                evidence.get("page"),
                evidence.get("chunk_id") or "",
                evidence.get("text") or "",
            ])

    workbook.save(file_path)


def _write_word(payload: dict[str, Any], file_path: Path) -> None:
    document = Document()
    document.add_heading("合同筛选结果", level=1)
    document.add_heading("筛选任务说明", level=2)
    document.add_paragraph(payload.get("prompt", ""))
    document.add_heading("筛选条件", level=2)
    document.add_paragraph(_conditions_text(payload) or "未提供结构化条件")
    document.add_heading("结果摘要", level=2)
    document.add_paragraph(f"命中合同数量：{len(payload.get('items') or [])}")
    document.add_paragraph(f"跳过文档：{payload.get('skipped') or {}}")
    document.add_heading("命中合同列表", level=2)
    for item in payload.get("items") or []:
        document.add_heading(item.get("name") or item.get("title") or "未命名合同", level=3)
        document.add_paragraph(f"文档 ID：{item.get('contract_id') or item.get('document_id') or ''}")
        document.add_paragraph(f"风险：{item.get('risk') or (item.get('meta') or {}).get('risk', '')}")
        document.add_paragraph(f"分数：{item.get('score') or (item.get('meta') or {}).get('score', '')}")
        document.add_paragraph(f"原因：{item.get('reason') or ''}")
        for evidence in item.get("evidence") or []:
            document.add_paragraph(
                f"{evidence.get('source') or '合同正文'} {evidence.get('ref') or ''}: {evidence.get('text') or ''}",
                style="List Bullet",
            )
    document.save(file_path)


def _conditions_text(payload: dict[str, Any]) -> str:
    strategy = payload.get("strategy") if isinstance(payload.get("strategy"), dict) else {}
    conditions = strategy.get("conditions") if isinstance(strategy.get("conditions"), list) else []
    return "；".join(condition.get("label", "") for condition in conditions if isinstance(condition, dict) and condition.get("label"))
